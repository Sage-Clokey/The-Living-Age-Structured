"""
Build WeeklyUpdateReport DOCX for Week 6.
Follows the template format from WeeklyUpdateReportTemplate.docx.

Usage: python paper/build_weekly_report_week6.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

TEMPLATE = Path(__file__).resolve().parent.parent / "WeeklyUpdateReportTemplate.docx"
OUT_DIR = Path(__file__).resolve().parent


def build_report(header, accomplishments, challenges, effort_rows, total_hours, out_name):
    doc = Document(str(TEMPLATE))

    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(header)
    run.font.size = Pt(11)

    while len(doc.paragraphs) > 4:
        p = doc.paragraphs[4]
        p._element.getparent().remove(p._element)

    for i, acc in enumerate(accomplishments, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Accomplishment {i}: {acc}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Challenges:")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()

    for i, ch in enumerate(challenges, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Challenge {i}: {ch}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    table = doc.tables[0]
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    for name, desc, hours, sig in effort_rows:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        row.cells[2].text = hours
        row.cells[3].text = sig

    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[2].text = total_hours

    out_path = OUT_DIR / out_name
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def build_week6():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 6 (May 4\u20139, 2026)"

    accomplishments = [
        "Attended SynBioBeta conference (May 4\u20136) in San Francisco. Spoke with Drew Endy "
        "about how biology design will decentralize things \u2014 directly validating the capstone\u2019s "
        "core thesis that living systems operate as decentralized economies and that engineering "
        "biology should work with that architecture, not against it.",

        "Had an extended conversation with Andrew Hessel about how living things are not machines "
        "and that living systems differ from computers in key ways. This reinforces the paper\u2019s "
        "central argument against the machine metaphor \u2014 the genome is not a program, the cell "
        "is not a factory, and engineering biology requires a fundamentally different paradigm.",

        "Discussed the themes of my writing with Karsten Temme (Pivot Bio) without revealing the "
        "specific work. Temme\u2019s company engineers nitrogen-fixing microbes for agriculture \u2014 "
        "a real-world example of working with living economies rather than replacing them. His "
        "perspective confirmed that the distributed design principles in the paper map to actual "
        "industry practice.",

        "Connected with a researcher who publishes papers in the synthetic biology space. "
        "Discussed the intersection of computational analysis and wet-lab validation, relevant "
        "to the capstone\u2019s experimental spread design.",

        "Built \u201cThe Price System of the Cell\u201d \u2014 a 4-panel synthesis figure that is the "
        "capstone\u2019s central visual. Panel A: three tiers of molecular prices (intracellular "
        "ratios, intercellular signals, mTOR as entrepreneur). Panel B: FBA shadow prices showing "
        "Menger\u2019s subjective value (same metabolite, different value in different conditions). "
        "Panel C: price discovery without a planner (metabolite pools oscillate then converge). "
        "Panel D: cancer breaks the price system (most mutated genes map to price components).",

        "Created comprehensive figure guide (figure_guide.md) mapping every claim in the capstone "
        "to a specific falsifiable visualization \u2014 what the figure should show if the thesis is "
        "correct and what the alternative would look like. Added price system layer to run_all.py "
        "pipeline.",
    ]

    challenges = [
        "Poster layout still pending. SynBioBeta consumed the first half of the week and the "
        "Price System figure work filled the second half. Poster is the top priority for Week 7.",

        "Need to translate the SynBioBeta conversations into concrete next steps. Drew Endy\u2019s "
        "perspective on decentralization and Andrew Hessel\u2019s critique of the machine metaphor "
        "both support the thesis, but should be cited carefully \u2014 these were informal "
        "conversations, not published positions.",

        "Paper finalization and sensitivity analysis on simulation parameters still outstanding. "
        "The figure guide now maps every claim to a visualization, which clarifies exactly what "
        "needs to be tightened before final submission.",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Attended SynBioBeta conference (May 4\u20136). Conversations with Drew Endy "
         "(decentralization in biology design), Andrew Hessel (living things are not machines), "
         "Karsten Temme (distributed design in agriculture), and others.",
         "12 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built Price System of the Cell 4-panel figure (generate_price_system_figure.py). "
         "Created individual panel exports. Added to run_all.py pipeline.",
         "3 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Created comprehensive figure guide mapping every claim to a falsifiable visualization. "
         "Built Week 6 progress report slides (7 slides with speaker notes).",
         "2 hrs", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "17 hrs",
                 "WeeklyUpdateReport_Week6.docx")


if __name__ == "__main__":
    build_week6()
