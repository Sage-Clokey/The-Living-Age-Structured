"""
Build Title and Abstract DOCX — intuition version.
Arial 11, single-spaced per BME129C_slides_051126.pdf guidelines.

Usage: python paper/build_intuition_abstract.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "intuition_abstract.docx"

FONT_NAME = "Arial"
FONT_SIZE = Pt(11)


def add_paragraph(doc, text, bold=False, italic=False, font_size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = 1.0
    pf.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    return p


def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_paragraph(doc,
        "Living systems as decentralized economies: biological networks coordinate "
        "through distributed prices, not central planning",
        bold=True, font_size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Author
    add_paragraph(doc, "Sage Clokey", font_size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc,
        "Department of Biomolecular Engineering and Bioinformatics, University of California, Santa Cruz",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc,
        "BME 129C: Design/Implement BME \u2014 Spring 2026",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc, "Advisor: R. Dubois",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    # Abstract heading
    add_paragraph(doc, "Abstract", bold=True, space_after=Pt(8))

    # Abstract body
    abstract = (
        "Living things are not machines. Living systems coordinate through "
        "distributed knowledge, not central planning. In biology there are no "
        "central banks, no legal tender laws \u2014 there is only voluntary trade. "
        "There are two types of order: centralized order requires uniformity, "
        "decentralized order requires ordered diversity. Life is decentralized, "
        "and life must die to become centralized. Life must be reduced to parts "
        "to be centralized, and in doing so it loses what makes it valuable."
        "\n\n"
        "The diversity of life is not random difference \u2014 it is based on local "
        "knowledge of time and place that no central planner can possess. "
        "Distributed nodes can plan but central planners cannot. The distributed "
        "knowledge is dictated by prices, the ratios between the voluntary "
        "exchange of anything between the nodes. Living order is organized by "
        "choice and state order is dictated by control."
        "\n\n"
        "Cells act by choice \u2014 they signal other cells, but they only react to "
        "those signals by choice, not coercion. There are cases where some coerce "
        "and that is called pathogens, viruses, and cancer. The immune system is "
        "the defense of the voluntary order."
        "\n\n"
        "The variance is not random. It is directed by distributed knowledge, "
        "just as the variance of the market is directed by prices. It is trial "
        "and error directed by knowledge \u2014 but the entrepreneurs do not die, "
        "they just pivot. The differences are a feature, not a bug \u2014 they are "
        "how adaptation is possible. Machines are reductionism to the point the "
        "parts are brittle; what was once strong is now weak. Once they become "
        "machines they are reduced to death."
        "\n\n"
        "Life is stories and bioinformatics is finding those stories so we can "
        "learn to become storytellers. Not to force form, but to help life grow. "
        "Evidence isn\u2019t to convince myself of what I already intuitively know "
        "\u2014 it is to convince others who don\u2019t have my intuition. That is the "
        "work. That is the spiral. That is the living age."
    )

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    run = p.add_run(abstract)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
