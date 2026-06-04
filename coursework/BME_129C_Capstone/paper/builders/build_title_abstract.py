"""
Build Title and Abstract DOCX for BME 129C Capstone.
Arial 11, single-spaced per BME129C_slides_051126.pdf guidelines.

Usage: python paper/build_title_abstract.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent.parent / "title_and_abstract_v3.docx"

FONT_NAME = "Arial"
FONT_SIZE = Pt(11)


def add_paragraph(doc, text, bold=False, italic=False, font_size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = 1.0
    pf.alignment = alignment
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_rich_paragraph(doc, segments, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = 1.0
    pf.alignment = alignment
    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = bold
        run.font.italic = italic
    return p


def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_paragraph(doc,
        "Living systems as decentralized economies: biological networks coordinate "
        "through distributed prices, not central planning",
        bold=True, font_size=Pt(12),
        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

    # Author
    add_paragraph(doc, "Sage Clokey", font_size=FONT_SIZE,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc,
        "Department of Biomolecular Engineering and Bioinformatics, University of California, Santa Cruz",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc,
        "BME 129C: Design/Implement BME \u2014 Spring 2026",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_paragraph(doc, "Advisor: R. Dubois",
        font_size=Pt(10), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(16))

    # Abstract heading
    add_paragraph(doc, "Abstract", bold=True, space_after=Pt(8))

    # Abstract body
    add_rich_paragraph(doc, [
        ("Living systems coordinate through distributed knowledge, not central "
         "planning. In biology there are no central banks, no legal tender laws "
         "\u2014 there is only voluntary trade. There are two types of order: "
         "centralized order requires uniformity, decentralized order requires "
         "ordered diversity. Life is decentralized, and the diversity of life is "
         "not random difference \u2014 it is based on local knowledge of time and "
         "place that no central planner can possess. The distributed knowledge is "
         "dictated by prices, the ratios between the voluntary exchange of anything "
         "between the nodes. Synthetic biology ignores this architecture, treating "
         "the cell as a machine to be commanded from above. This study tests whether "
         "that machine metaphor misrepresents how living systems actually organize."
         "\n\n"
         "We analyze biological networks, metabolic coordination, single-cell "
         "specialization, and cross-species compatibility and find the same pattern "
         "at every level: decentralized order with no master node, no master cell, "
         "and no central planner. Biological networks tolerate massive node removal "
         "while centralized architectures collapse. Metabolic pathways modeled as "
         "economic agents reach equilibrium through local feedback alone and "
         "outperform centralized allocation under perturbation. Cross-species gene "
         "transferability follows patterns of voluntary exchange, with each organism "
         "contributing unique comparative advantage. Cells act by choice and signal "
         "other cells, but they react to those signals by choice. Human immune cells "
         "confirm this: eight specialized cell types differentiating without a master "
         "cell, communicating without a gatekeeper, and surviving any single cell "
         "type removal."
         "\n\n"
         "The variance is not random. It is directed by distributed knowledge, just "
         "as the variance of the market is directed by prices. Living things are not "
         "machines \u2014 once they become machines they are reduced to death. If DNA is "
         "a language, the cell speaks it and can write in that language. "
         "Bioinformatics is the study of this distributed knowledge in living "
         "systems. Build economies, not machines.", False, False),
    ])

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
