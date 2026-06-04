"""
Generate Methods Section as a .docx file.
Arial 11pt, single-spaced, subsection headers.

Usage: python paper/methods_section.docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 1.0

# Heading style
for level in [1, 2, 3]:
    h_style = doc.styles[f"Heading {level}"]
    h_style.font.name = "Arial"
    h_style.font.size = Pt(12) if level == 1 else Pt(11)
    h_style.font.bold = True
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    h_style.paragraph_format.space_before = Pt(12)
    h_style.paragraph_format.space_after = Pt(4)
    h_style.paragraph_format.line_spacing = 1.0


def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


# ====================================================================
# TITLE
# ====================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Methods")
run.bold = True
run.font.size = Pt(14)
run.font.name = "Arial"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run(
    "Living Systems as Decentralized Economies\n"
    "BME 129C Capstone — Sage Clokey — Spring 2026\n"
    "Advisor: R. Dubois, UC Santa Cruz"
)
run2.font.size = Pt(10)
run2.font.name = "Arial"
run2.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()  # spacer

# ====================================================================
# 1. Databases and Software
# ====================================================================
doc.add_heading("1. Databases, Software, and Data Sources", level=1)

add_body(
    "All analyses were implemented in Python 3.12. The following databases and software "
    "libraries were used. Network construction and graph analysis used NetworkX 3.6 "
    "(Hagberg et al., 2008). Single-cell RNA sequencing analysis used Scanpy 1.12 "
    "(Wolf et al., 2018). Genome-scale metabolic modeling used COBRApy 0.31 "
    "(Ebrahim et al., 2013). Power-law fitting used the powerlaw library 2.0 "
    "(Alstott et al., 2014; method of Clauset, Shalizi, & Newman, 2009). "
    "Statistical analysis used SciPy 1.17 and NumPy 2.4. "
    "Visualization used Matplotlib 3.10 and python-pptx 1.0. "
    "The complete analysis pipeline is executable via a single command "
    "(python run_all.py) and is version-controlled on GitHub."
)

# --- Table: Databases ---
doc.add_heading("Table 1. Databases and Accession Numbers", level=2)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, text in enumerate(["Database", "URL / Accession", "Data Retrieved", "Layer"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.name = "Arial"

rows_data = [
    ["RegulonDB v11", "regulondb.ccg.unam.mx", "E. coli K-12 TF-gene interactions (282 nodes, 308 edges)", "1"],
    ["KEGG REST API", "rest.kegg.jp\neco00010, eco00020, eco00190, eco00230, eco00240, eco00061, eco00260, eco00620", "E. coli and S. cerevisiae metabolic pathway networks", "1"],
    ["STRING v12", "string-db.org\nTaxon 511145 (E. coli), 4932 (yeast)", "Protein-protein interaction networks (score >= 0.7)", "1"],
    ["10x Genomics / Scanpy", "scanpy.datasets.pbmc3k()\nGEO: GSE115189 (related)", "Human PBMC scRNA-seq (2,638 cells, 13,714 genes, 8 cell types)", "1b"],
    ["BiGG Models (iML1515)", "bigg.ucsd.edu/models/iML1515\nMonk et al. 2017, Nat Biotechnol 35:904", "E. coli genome-scale metabolic model (2,712 rxns, 1,877 mets, 1,516 genes)", "2"],
    ["Keio Collection", "Baba et al. 2006, Mol Syst Biol 2:2006.0008", "Systematic single-gene knockout growth phenotypes (3,985 genes)", "2"],
    ["Kazusa Codon Usage DB", "kazusa.or.jp/codon/", "RSCU tables for 8 organisms across 4 kingdoms", "3"],
    ["Adaptive Genome Design System", "github.com/Sage-Clokey/adaptive-Automation", "13 pathway profiles, codon compatibility engine, capability map", "2, 3"],
]

for rd in rows_data:
    row = table.add_row()
    for i, text in enumerate(rd):
        row.cells[i].text = text
        for p in row.cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
                r.font.name = "Arial"

doc.add_paragraph()

# ====================================================================
# 2. Layer 1: Network Topology
# ====================================================================
doc.add_heading("2. Biological Network Construction and Topology Analysis (Layer 1)", level=1)

add_body(
    "Gene regulatory network. The E. coli K-12 gene regulatory network was constructed from "
    "curated transcription factor (TF) to target gene interactions representing the RegulonDB v11 "
    "reference dataset (Santos-Zavaleta et al., 2019). The network encodes 282 nodes (25 major "
    "transcription factors and their validated targets) and 308 directed edges. Key global regulators "
    "include CRP (43 targets), FNR (20 targets), ArcA (15 targets), and LexA (14 targets)."
)

add_body(
    "Metabolic networks. Metabolic pathway networks for E. coli (organism code: eco) and "
    "S. cerevisiae (sce) were retrieved from the KEGG REST API (Kanehisa & Goto, 2000). "
    "Eight core E. coli pathways were included: glycolysis (eco00010), TCA cycle (eco00020), "
    "oxidative phosphorylation (eco00190), purine metabolism (eco00230), pyrimidine metabolism "
    "(eco00240), fatty acid biosynthesis (eco00061), glycine/serine/threonine metabolism "
    "(eco00260), and pyruvate metabolism (eco00620). Each pathway was encoded as a directed "
    "graph with genes and compounds as nodes and edges representing co-participation in "
    "metabolic reactions."
)

add_body(
    "Protein-protein interaction networks. High-confidence PPI networks for E. coli "
    "(taxonomy ID 511145) and S. cerevisiae (taxonomy ID 4932) were retrieved from the STRING "
    "database v12 (Szklarczyk et al., 2023) using 40 seed genes per organism. Only interactions "
    "with combined confidence score >= 0.7 were retained."
)

add_body(
    "Reference networks. Five synthetic comparison networks were constructed using NetworkX to "
    "represent alternative architectures: (1) star graph (pure centralization), (2) Erdos-Renyi "
    "random graph, (3) regular lattice, (4) hub-and-spoke (10 hubs), and (5) Barabasi-Albert "
    "scale-free model. All reference networks were matched to the largest biological network "
    "in node count."
)

doc.add_heading("Topology Metrics", level=2)

add_body(
    "Degree distribution and power-law fitting. The degree sequence of each network was fit to "
    "a discrete power-law distribution using the method of Clauset, Shalizi, and Newman (2009). "
    "Scale-free status was assessed by comparing power-law fit to an exponential alternative "
    "using log-likelihood ratio R with significance at p < 0.05."
)

add_body(
    "Betweenness centrality and Gini coefficient. Betweenness centrality was computed for all "
    "nodes using Brandes' algorithm (2001). The Gini coefficient of the betweenness distribution "
    "was computed as a measure of centralization inequality."
)

add_body(
    "Robustness analysis. Network robustness was assessed by iterative node removal under two "
    "strategies following Albert et al. (2000): (a) random removal, where nodes are removed "
    "uniformly at random, and (b) targeted attack, where the highest-degree node is removed at "
    "each step with degree recalculated after each removal. The robustness threshold was defined "
    "as the fraction of nodes removed before the giant component drops below 50%."
)

# ====================================================================
# 3. Motif Analysis
# ====================================================================
doc.add_heading("3. Network Motif Analysis (Alon)", level=1)

add_body(
    "Network motifs were identified using the method of Alon et al. (2002). A triadic census "
    "was performed on the E. coli GRN, counting all 16 possible three-node subgraph isomorphism "
    "classes (triad types). To determine statistical significance, 1,000 degree-preserving "
    "randomized networks were generated by repeated edge swapping (Maslov & Sneppen, 2002), "
    "preserving the in-degree and out-degree of every node while destroying higher-order "
    "structure. For each triad type, the Z-score was computed as (real count - mean random count) "
    "/ standard deviation of random counts. Triad types with Z > 2.0 were classified as motifs "
    "(over-represented); Z < -2.0 as anti-motifs (suppressed). The feed-forward loop (triad "
    "type 030T) was specifically analyzed as a candidate Hayekian price signal motif."
)

# ====================================================================
# 4. Self-Regulation (WBPA)
# ====================================================================
doc.add_heading("4. Self-Regulation Analysis (WBPA)", level=1)

add_body(
    "The self-regulation mechanism was analyzed following the Weighted Betweenness Preferential "
    "Attachment (WBPA) model of Topirceanu et al. (2018). For each network, the Spearman rank "
    "correlation between node degree and betweenness centrality was computed. A low correlation "
    "(rho < 0.8) indicates that high-degree nodes do not monopolize information flow — the "
    "network self-regulates against centralization. The Gini ratio (betweenness Gini / degree "
    "Gini) was computed as a self-regulation index; values below 1.0 indicate that betweenness "
    "is distributed more equally than degree alone would predict."
)

add_body(
    "Hub erosion test. To test the WBPA mechanism directly, edges were artificially added to "
    "the top 5 highest-degree nodes (10 edges per hub), and the change in betweenness centrality "
    "was measured. In self-regulating networks, connecting to a hub should decrease its "
    "betweenness (the network routes around it); in centralizing networks, betweenness should "
    "increase. Three growth models were compared: degree-preferential attachment (Barabasi-Albert), "
    "betweenness-preferential attachment (WBPA), and random attachment, each grown to 300 nodes."
)

# ====================================================================
# 5. Single-Cell Economy
# ====================================================================
doc.add_heading("5. Single-Cell Economy Analysis (Layer 1b)", level=1)

add_body(
    "Data. The human PBMC dataset (pbmc3k) was obtained from 10x Genomics via the Scanpy "
    "built-in data loader (Wolf et al., 2018). The dataset contains 2,638 cells across 8 cell "
    "types: B cells, CD4 T cells, CD8 T cells, CD14+ monocytes, FCGR3A+ monocytes, dendritic "
    "cells, megakaryocytes, and NK cells. The raw count matrix (13,714 genes) was preprocessed "
    "with standard quality control (minimum 200 genes per cell, minimum 3 cells per gene), "
    "total-count normalization to 10,000 reads per cell, and log1p transformation."
)

add_body(
    "Specialization scores. For each cell type, Shannon entropy of the mean expression vector "
    "was computed and normalized by log(number of expressed genes). Lower normalized entropy "
    "indicates higher specialization — the cell type concentrates its transcriptional output "
    "on a narrow gene program."
)

add_body(
    "Cell-cell communication network. Intercellular communication was inferred from 30 curated "
    "ligand-receptor pairs relevant to immune signaling, drawn from Ramilowski et al. (2015) "
    "and CellChat (Jin et al., 2021). A pair was scored as active between cell types A and B "
    "if both the ligand (in A) and receptor (in B) exceeded a mean expression threshold of 0.1 "
    "in log-normalized space. Betweenness Gini of the resulting network was computed to assess "
    "communication centralization. Robustness was tested by removing each cell type and measuring "
    "the fraction of communication edges surviving."
)

add_body(
    "Comparative advantage. Fold change of mean expression per gene per cell type versus all "
    "other cell types was computed. The top 10 enriched genes per cell type were identified as "
    "that cell type's comparative advantage."
)

# ====================================================================
# 6. Economic Simulation (Layer 2)
# ====================================================================
doc.add_heading("6. Agent-Based Economic Simulation (Layer 2)", level=1)

add_body(
    "Pathway agents. Thirteen metabolic pathway profiles from the Adaptive Genome Design System "
    "were modeled as economic agents. Each agent has defined substrates consumed (demand), "
    "products generated (supply), ATP cost (currency), and signal inputs/outputs (price signals). "
    "Pathways span four kingdoms: bacterial cellulose synthesis (Komagataeibacter), fungal chitin "
    "synthesis (Ganoderma), coral biomineralization (Acropora), spider silk (Trichonephila), "
    "firefly bioluminescence (Photinus), heat shock (universal), HIF1a hypoxia (metazoan), "
    "Shh patterning (vertebrate), WUS/CLV3 meristem (Arabidopsis), aquaporin (universal), "
    "Piwi regeneration (planarian), MMP9 remodeling (vertebrate), and sea urchin biomineralization."
)

add_body(
    "Three allocation regimes were compared. (1) Distributed (biological): each agent adjusts "
    "its own production rate based on local metabolite concentrations — product feedback (reduce "
    "if oversupplied), substrate feedback (reduce if scarce), and ATP feedback (reduce if "
    "depleted). No agent has information about any other agent. (2) Centralized (naive): a global "
    "allocator ranks agents by ATP efficiency and assigns rates proportional to rank. (3) "
    "Centralized-smart: a global allocator with real-time metabolite pool access re-optimizes "
    "every step based on current scarcity — the 'socialist planning board with a supercomputer' "
    "(Mises' strongest opponent). All regimes were run for 200 time steps with ATP regeneration "
    "of 2.0 per agent per step and identical external metabolite supply."
)

add_body(
    "Perturbation suite. Four perturbation types were applied at step 100: (a) substrate shock — "
    "80% reduction in a key substrate supply, testing Hayek's knowledge problem; (b) ATP crisis — "
    "50% reduction in ATP regeneration for 30 steps, testing Mises' calculation problem; (c) "
    "demand spike — 3x demand for one product, testing Kirzner's entrepreneurial discovery; (d) "
    "novel opportunity — introduction of a new metabolite resource, testing Kirzner's "
    "entrepreneurial alertness. Recovery time, trough depth, and GDP gain were measured for all "
    "three regimes under each perturbation."
)

# ====================================================================
# 7. FBA Analysis (Layer 2)
# ====================================================================
doc.add_heading("7. Flux Balance Analysis (Layer 2)", level=1)

add_body(
    "Genome-scale metabolic model. The E. coli K-12 MG1655 genome-scale metabolic model iML1515 "
    "(Monk et al., 2017; 2,712 reactions, 1,877 metabolites, 1,516 genes) was loaded via "
    "COBRApy 0.31 (Ebrahim et al., 2013) from the BiGG Models database. Flux Balance Analysis "
    "(FBA) optimizes the linear program max(c^T * v) subject to S*v = 0 and v_lb <= v <= v_ub, "
    "where S is the stoichiometric matrix, v is the flux vector, and c is the objective "
    "(biomass production). This represents the strongest possible central planner: an omniscient "
    "allocator with perfect knowledge of all 2,712 reaction constraints solving simultaneously "
    "for optimal growth."
)

add_body(
    "Gene knockout analysis. Single-gene knockouts were simulated for 51 genes (35 essential "
    "and 16 non-essential) using COBRApy's gene knockout function, which properly handles "
    "gene-protein-reaction (GPR) associations including AND/OR logic. Gene names were resolved "
    "to iML1515 b-numbers (e.g., pgk -> b2926) via the model's gene name annotations. "
    "A gene was classified as FBA-essential if knockout growth fell below 5% of wildtype. "
    "FBA predictions were compared to published Keio collection growth phenotypes (Baba et al., "
    "2006) to compute accuracy, sensitivity, and specificity of the omniscient planner."
)

add_body(
    "Shadow price analysis. Shadow prices (dual variables of the LP) were extracted from the "
    "FBA solution. These represent the marginal growth value of each metabolite constraint — "
    "the Hayekian price signal computed internally by the planner."
)

add_body(
    "Environmental perturbation. Three perturbations were applied to the iML1515 model: "
    "(a) carbon source switch (glucose -> acetate, EX_glc__D_e set to 0, EX_ac_e set to -10 "
    "mmol/gDW/h); (b) anaerobic growth (EX_o2_e set to 0); (c) nitrogen limitation (EX_nh4_e "
    "set to -1 mmol/gDW/h). For each perturbation, FBA was re-solved and the predicted growth "
    "ratio relative to wildtype was compared to known biological responses (diauxic shift, "
    "ArcAB/FNR regulon activation, NtrBC nitrogen sensing)."
)

add_body(
    "Conditionally essential genes. Ten genes known to be essential in minimal media but "
    "non-essential in rich media (trpA, trpB, hisB, proA, serA, cysE, metA, lysA, pheA, tyrA) "
    "were analyzed to test Menger's subjective value: essentiality is not an intrinsic property "
    "of the gene but depends on environmental context."
)

# ====================================================================
# 8. Cross-Species Trade (Layer 3)
# ====================================================================
doc.add_heading("8. Cross-Species Trade Network (Layer 3)", level=1)

add_body(
    "Organisms. Five to eight organisms spanning four kingdoms were modeled: Komagataeibacter "
    "xylinus (prokaryote), E. coli K-12 (prokaryote), S. cerevisiae (fungus), Ganoderma "
    "lucidum (fungus), Arabidopsis thaliana (plant), Homo sapiens (animal), Acropora millepora "
    "(coral), and Ambystoma mexicanum (axolotl)."
)

add_body(
    "Codon usage distance. Relative Synonymous Codon Usage (RSCU) tables were obtained from "
    "the Kazusa Codon Usage Database (Nakamura et al., 2000) and the Adaptive Genome Design "
    "System. Pairwise codon distance was computed as the Euclidean distance between 61-dimensional "
    "RSCU vectors."
)

add_body(
    "Trade cost. Total trade cost was computed as: C_trade(A,B) = 0.5 * d_codon(A,B) + 0.3 * "
    "R_regulatory(A,B) + 0.2 * B_baseline, where d_codon is codon distance, R_regulatory is a "
    "regulatory barrier (0.3 for cross-kingdom, 0.1 for same-kingdom), and B_baseline is a "
    "constant baseline cost."
)

add_body(
    "Voluntary exchange validation. FBA predictions were validated against 18 published "
    "heterologous expression experiments (GFP in E. coli, human insulin in E. coli, spider "
    "silk fibroin in yeast, etc.) with documented success scores. Trade cost was correlated "
    "with expression success using Spearman rank correlation. Trade blocs were detected using "
    "Louvain community detection on the trade network."
)

# ====================================================================
# 9. Data Availability
# ====================================================================
doc.add_heading("9. Data Availability", level=1)

add_body(
    "All code, analysis pipelines, and cached data for this study are available on GitHub at "
    "https://github.com/Sage-Clokey/Living-works-by-the-word in the BME_129C_Capstone/ "
    "directory. The complete analysis can be reproduced by running 'python run_all.py' from the "
    "project root. Individual layers can be run separately (e.g., 'python run_all.py --layer 2f' "
    "for FBA analysis). The --quick flag skips external API calls and uses cached/built-in data."
)

add_body(
    "Raw data sources are publicly available: RegulonDB (regulondb.ccg.unam.mx), KEGG "
    "(rest.kegg.jp), STRING (string-db.org), 10x Genomics pbmc3k dataset (via Scanpy), BiGG "
    "Models iML1515 (bigg.ucsd.edu), and Kazusa Codon Usage Database (kazusa.or.jp/codon/). "
    "The Adaptive Genome Design System, which provides pathway profiles and codon compatibility "
    "infrastructure, is available at https://github.com/Sage-Clokey/adaptive-Automation."
)

add_body(
    "Generated figures are saved to paper/figures/ (combined multi-panel) and "
    "paper/figures/individual/ (one panel per PNG). The presentation is generated by "
    "'python paper/generate_presentation.py' and saved as paper/capstone_presentation.pptx."
)

# ====================================================================
# Save
# ====================================================================
out_path = Path(__file__).resolve().parent.parent / "methods_section.docx"
doc.save(str(out_path))
print(f"Methods section saved to: {out_path}")
