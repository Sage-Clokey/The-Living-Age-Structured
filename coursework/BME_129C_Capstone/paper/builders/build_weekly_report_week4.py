"""
Build WeeklyUpdateReport_Week4.docx from the template.
Week 4: April 14-18, 2026
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy

TEMPLATE = Path(__file__).resolve().parent.parent / "WeeklyUpdateReportTemplate.docx"
OUT = Path(__file__).resolve().parent.parent / "WeeklyUpdateReport_Week4.docx"


def build():
    doc = Document(str(TEMPLATE))

    # --- Header ---
    doc.paragraphs[0].text = "Student: Sage Clokey (Individual Project) — Week 4 (April 14–18, 2026)"

    # --- Accomplishments ---
    accomplishments = [
        (
            "Accomplishment 1: Integrated Hayek's The Fatal Conceit (1988) across the capstone paper "
            "and slide deck. The core argument — 'The curious task of economics is to demonstrate to men "
            "how little they really know about what they imagine they can design' — now frames the "
            "introduction, discussion, and conclusion. Applied directly to synthetic biology: redesigning "
            "evolved order destroys distributed information."
        ),
        (
            "Accomplishment 2: Wrote Section 4.2 — Design Principles: Engineering Economies, Not Machines. "
            "Two concrete architectures for sagent engineering: (1) single organism designed FOR evolution "
            "(feedback promoters, shadow-price-guided expression, codon harmonization), and (2) microbial "
            "consortium with division of labor mirroring PBMC cell economy data."
        ),
        (
            "Accomplishment 3: Designed Section 4.7 — Experimental Spread with 6 violacein pathway "
            "experiments (vioABCDE from Chromobacterium violaceum). Each experiment tests one Austrian "
            "economic prediction with wet-lab data: (1) fixed vs feedback promoters (Hayek), "
            "(2) star vs distributed control (Mises), (3) single strain vs consortium (Menger), "
            "(4) codon optimization vs harmonization (Layer 3), (5) adaptive evolution (Kirzner), "
            "(6) perturbation gauntlet (Rothbard). All measured by violacein absorbance at 575nm."
        ),
        (
            "Accomplishment 4: Overhauled all figure code with a new visualization principle — show every "
            "individual data point. Bar charts replaced with semi-transparent bars + scatter overlays, "
            "robustness curves show individual trial lines behind averages, hub erosion uses per-hub "
            "jitter plots with median lines. Updated 6 analysis files and regenerated all 21 figures."
        ),
        (
            "Accomplishment 5: Fixed price_signals bug in single_cell_economy.py and regenerated the full "
            "pipeline through run_all.py. All layers now produce consistent, reproducible output."
        ),
    ]

    # Clear template placeholder accomplishments (paragraphs 4-5)
    # Rebuild accomplishments section
    doc.paragraphs[4].text = accomplishments[0]
    doc.paragraphs[5].text = ""

    # Insert additional accomplishment paragraphs
    for i, acc in enumerate(accomplishments[1:], start=1):
        p = doc.paragraphs[5].insert_paragraph_after(acc) if i == 1 else p.insert_paragraph_after(acc)
        p.style = doc.paragraphs[4].style
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # Find and update challenges section
    # We need to find the challenges paragraphs and update them
    challenge_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Challenges"):
            challenge_idx = i
            break

    challenges = [
        (
            "Challenge 1: Experimental spread (6 violacein experiments) needs advisor validation for "
            "feasibility within available wet-lab resources. Plan to prioritize experiments 1-3 if "
            "resources are limited. Meeting with Dubois and Lowe needed."
        ),
        (
            "Challenge 2: Statistical validation not yet integrated — need Mann-Whitney U for distributed "
            "vs centralized GDP distributions, bootstrap CIs on robustness metrics, and KS test for "
            "power-law degree distribution fit. Planned for Week 5."
        ),
    ]

    if challenge_idx is not None:
        # Update existing challenge placeholders
        for i, ch in enumerate(challenges):
            idx = challenge_idx + 2 + i  # +2 to skip "Challenges:" header and blank line
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].text = ch

    # --- Table ---
    table = doc.tables[0]
    rows_data = [
        ("Sage Clokey", "Integrated Fatal Conceit (Hayek 1988) into paper introduction, discussion, and conclusion. Wrote Design Principles section (Section 4.2) with two architectures for sagent engineering.", "3 hrs", "Sage Clokey"),
        ("Sage Clokey", "Designed experimental spread (Section 4.7): 6 violacein experiments mapping Austrian predictions to wet-lab tests. Researched vioABCDE pathway and measurement protocols.", "3 hrs", "Sage Clokey"),
        ("Sage Clokey", "Overhauled all figure code across 6 analysis files. New data-point-forward visualization principle. Regenerated all 21 figures through run_all.py pipeline.", "3 hrs", "Sage Clokey"),
        ("Sage Clokey", "Fixed price_signals bug. Built Week 4 progress report slides (PPTX). Paper editing and polish.", "2 hrs", "Sage Clokey"),
        ("TOTAL", "", "11 hrs", ""),
    ]

    # Clear existing data rows and add new ones
    # Remove extra template rows (keep header)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
