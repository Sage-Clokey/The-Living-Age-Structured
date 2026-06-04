#!/usr/bin/env python3
"""
build_paper_docx.py
Converts capstone_paper.md into a formatted Word document (.docx).

Usage:
    python build_paper_docx.py
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = Path(__file__).resolve().parent
INPUT_MD = SCRIPT_DIR / "capstone_paper.md"
OUTPUT_DOCX = SCRIPT_DIR / "capstone_paper_week4.docx"
FIGURES_DIR = SCRIPT_DIR / "figures"

# ---------------------------------------------------------------------------
# Formatting constants
# ---------------------------------------------------------------------------
FONT_BODY = "Times New Roman"
FONT_MONO = "Consolas"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(14)
H2_SIZE = Pt(13)
H3_SIZE = Pt(12)
H4_SIZE = Pt(11)
QUOTE_SIZE = Pt(10)
CODE_SIZE = Pt(9)
REF_SIZE = Pt(10)


# ---------------------------------------------------------------------------
# Inline formatting helpers
# ---------------------------------------------------------------------------

def add_inline_runs(paragraph, text, base_bold=False, base_italic=False,
                    base_size=BODY_SIZE, base_font=FONT_BODY):
    """Parse inline markdown (bold, italic, code, bold-italic) and add runs."""
    # Pattern to match inline formatting tokens
    # Order matters: bold-italic first, then bold, italic, code
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # bold italic
        r'|(\*\*(.+?)\*\*)'      # bold
        r'|(\*(.+?)\*)'          # italic
        r'|(`(.+?)`)'            # inline code
    )
    pos = 0
    for m in pattern.finditer(text):
        # Add text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = base_size
            run.font.name = base_font
            run.bold = base_bold
            run.italic = base_italic

        if m.group(2):  # bold italic ***...***
            run = paragraph.add_run(m.group(2))
            run.font.size = base_size
            run.font.name = base_font
            run.bold = True
            run.italic = True
        elif m.group(4):  # bold **...**
            run = paragraph.add_run(m.group(4))
            run.font.size = base_size
            run.font.name = base_font
            run.bold = True
            run.italic = base_italic
        elif m.group(6):  # italic *...*
            run = paragraph.add_run(m.group(6))
            run.font.size = base_size
            run.font.name = base_font
            run.bold = base_bold
            run.italic = True
        elif m.group(8):  # code `...`
            run = paragraph.add_run(m.group(8))
            run.font.size = CODE_SIZE
            run.font.name = FONT_MONO
            run.bold = False
            run.italic = False

        pos = m.end()

    # Remaining text after last match
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = base_size
        run.font.name = base_font
        run.bold = base_bold
        run.italic = base_italic


def set_paragraph_format(paragraph, size=BODY_SIZE, bold=False, italic=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=Pt(0), space_after=Pt(6),
                         font_name=FONT_BODY, first_line_indent=None,
                         left_indent=None, hanging_indent=None,
                         line_spacing=1.0):
    """Set paragraph-level formatting."""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if hanging_indent is not None:
        # Hanging indent: set left_indent and negative first_line_indent
        if left_indent is None:
            pf.left_indent = hanging_indent
        pf.first_line_indent = -hanging_indent


def add_page_number(doc):
    """Add page numbers to the footer."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # PAGE field
    run = p.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld_char_begin)
    run2 = p.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instr)
    run3 = p.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld_char_end)
    for r in [run, run2, run3]:
        r.font.size = Pt(9)
        r.font.name = FONT_BODY


def add_thin_line(doc):
    """Add a thin horizontal line (border-bottom on an empty paragraph)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_figure(doc, img_path, caption_text):
    """Embed a figure with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)

    resolved = FIGURES_DIR / img_path if not os.path.isabs(img_path) else Path(img_path)
    if not resolved.exists():
        # Try stripping leading "figures/"
        alt = FIGURES_DIR / Path(img_path).name
        if alt.exists():
            resolved = alt

    if resolved.exists() and resolved.suffix.lower() in ('.png', '.jpg', '.jpeg', '.bmp', '.gif'):
        run = p.add_run()
        run.add_picture(str(resolved), width=Inches(5.5))
    else:
        run = p.add_run(f"[Image not found: {img_path}]")
        run.font.size = BODY_SIZE
        run.font.name = FONT_BODY
        run.italic = True

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    add_inline_runs(cap, caption_text, base_italic=True, base_size=Pt(10))


def add_table(doc, table_lines):
    """Parse markdown table lines and insert a Word table."""
    # Split cells
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)

    if len(rows) < 2:
        return

    # Check for separator row (---|---|...)
    sep_idx = None
    for i, row in enumerate(rows):
        if all(re.match(r'^[-:]+$', c) for c in row if c):
            sep_idx = i
            break

    data_rows = [r for i, r in enumerate(rows) if i != sep_idx]
    if not data_rows:
        return

    n_cols = max(len(r) for r in data_rows)
    # Normalize column counts
    for r in data_rows:
        while len(r) < n_cols:
            r.append('')

    tbl = doc.add_table(rows=len(data_rows), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for i, row_data in enumerate(data_rows):
        for j, cell_text in enumerate(row_data):
            cell = tbl.cell(i, j)
            # Clear default paragraph
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text, base_size=Pt(9), base_font=FONT_BODY)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Header row: bold
            if i == 0:
                for run in p.runs:
                    run.bold = True

            # Shade header row
            if i == 0:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
                )
                cell._element.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def build_docx():
    md_text = INPUT_MD.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = BODY_SIZE
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(6)

    add_page_number(doc)

    i = 0
    in_references = False
    title_done = False
    author_block_done = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- Title (# ...) ---
        if not title_done and stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(title_text)
            run.font.size = TITLE_SIZE
            run.font.name = FONT_BODY
            run.bold = True
            title_done = True
            i += 1
            continue

        # --- Author block (lines 2-5 after title, before first ##) ---
        if title_done and not author_block_done and not stripped.startswith('#'):
            # Gather author/affiliation lines until we hit a heading, ---, or >
            if stripped == '':
                i += 1
                continue
            if stripped.startswith('---') or stripped.startswith('>') or stripped.startswith('##'):
                author_block_done = True
                continue  # Don't advance i — process this line normally

            # It's an author/affiliation line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            add_inline_runs(p, stripped, base_size=BODY_SIZE, base_font=FONT_BODY)
            i += 1
            continue

        if not author_block_done and title_done:
            author_block_done = True

        # --- Horizontal rule ---
        if stripped == '---' or stripped == '***' or stripped == '___':
            add_thin_line(doc)
            i += 1
            continue

        # --- Section headers ---
        h4_match = re.match(r'^####\s+(.*)', stripped)
        h3_match = re.match(r'^###\s+(.*)', stripped)
        h2_match = re.match(r'^##\s+(.*)', stripped)

        if h2_match:
            header_text = h2_match.group(1).strip()
            if header_text.lower() == 'references':
                in_references = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(header_text)
            run.font.size = H2_SIZE
            run.font.name = FONT_BODY
            run.bold = True
            i += 1
            continue

        if h3_match and not h4_match:
            header_text = h3_match.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(header_text)
            run.font.size = H3_SIZE
            run.font.name = FONT_BODY
            run.bold = True
            run.italic = True
            i += 1
            continue

        if h4_match:
            header_text = h4_match.group(1).strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(header_text)
            run.font.size = H4_SIZE
            run.font.name = FONT_BODY
            run.bold = True
            run.italic = True
            i += 1
            continue

        # --- Figure images ---
        fig_match = re.match(r'^!\[(.+?)\]\((.+?)\)', stripped)
        if fig_match:
            caption = fig_match.group(1)
            img_path = fig_match.group(2)
            add_figure(doc, img_path, caption)
            i += 1
            continue

        # --- Figure caption lines (bold text after a figure image) ---
        # e.g. **Figure 1. ...** — already handled by inline formatting

        # --- Block quotes ---
        if stripped.startswith('>'):
            # Collect consecutive blockquote lines
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qtext = lines[i].strip()
                # Remove leading > and optional space
                qtext = re.sub(r'^>\s?', '', qtext)
                quote_lines.append(qtext)
                i += 1
            quote_text = ' '.join(quote_lines)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.0
            add_inline_runs(p, quote_text, base_italic=True, base_size=QUOTE_SIZE)
            continue

        # --- Tables ---
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        # --- Bullet lists ---
        bullet_match = re.match(r'^(\s*)-\s+(.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1))
            bullet_text = bullet_match.group(2).strip()

            # Collect continuation lines (non-empty, non-bullet, indented more)
            i += 1
            while i < len(lines):
                next_line = lines[i]
                next_stripped = next_line.strip()
                # continuation: indented text that is not a new bullet and not empty
                if next_stripped and not next_stripped.startswith('-') and not next_stripped.startswith('#') \
                        and not next_stripped.startswith('|') and not next_stripped.startswith('>') \
                        and not next_stripped.startswith('!') and len(next_line) - len(next_line.lstrip()) > 0:
                    bullet_text += ' ' + next_stripped
                    i += 1
                else:
                    break

            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            if indent_level >= 2:
                # Sub-bullet
                p.paragraph_format.left_indent = Inches(0.75)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                marker = '\u2013 '  # en-dash for sub-bullets
            else:
                # Top-level bullet
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                marker = '\u2022 '  # bullet character

            run = p.add_run(marker)
            run.font.size = BODY_SIZE
            run.font.name = FONT_BODY
            add_inline_runs(p, bullet_text, base_size=BODY_SIZE)
            continue

        # --- Empty lines ---
        if stripped == '':
            i += 1
            continue

        # --- References section: hanging indent ---
        if in_references:
            # Collect full reference (may span multiple lines)
            ref_text = stripped
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#'):
                # Check if next line is a continuation (doesn't start with a
                # new reference — references start with a capital letter or author name)
                next_s = lines[i].strip()
                # If next line looks like it starts a new reference (has a year pattern),
                # break; otherwise it's a continuation
                if re.match(r'^[A-Z].*\(\d{4}\)', next_s):
                    break
                ref_text += ' ' + next_s
                i += 1

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, ref_text, base_size=REF_SIZE)
            continue

        # --- Normal body paragraph ---
        para_text = stripped
        i += 1
        # Collect continuation lines for the paragraph
        while i < len(lines):
            next_line = lines[i]
            next_stripped = next_line.strip()
            if next_stripped == '':
                break
            if next_stripped.startswith('#') or next_stripped.startswith('>') \
                    or next_stripped.startswith('|') or next_stripped.startswith('---') \
                    or next_stripped.startswith('***') or next_stripped.startswith('___') \
                    or next_stripped.startswith('- ') or next_stripped.startswith('  -') \
                    or re.match(r'^!\[', next_stripped):
                break
            para_text += ' ' + next_stripped
            i += 1

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        add_inline_runs(p, para_text, base_size=BODY_SIZE)

    # --- Save ---
    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    build_docx()
