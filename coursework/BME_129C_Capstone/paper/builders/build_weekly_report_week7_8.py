"""
Build WeeklyUpdateReport DOCX files for Weeks 7 and 8.
Follows the template format from WeeklyUpdateReportTemplate.docx.

Usage: python paper/builders/build_weekly_report_week7_8.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

TEMPLATE = Path(__file__).resolve().parent.parent.parent / "WeeklyUpdateReportTemplate.docx"
OUT_DIR = Path(__file__).resolve().parent.parent / "weekly_reports"


def build_report(header, accomplishments, challenges, effort_rows, total_hours, out_name):
    doc = Document(str(TEMPLATE))

    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(header)
    run.font.size = Pt(11)

    while len(doc.paragraphs) > 4:
        p = doc.paragraphs[4]
        p._element.getparent().remove(p._element)

    for i, acc in enumerate(accomplishments, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Accomplishment {i}: {acc}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Challenges:")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()

    for i, ch in enumerate(challenges, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Challenge {i}: {ch}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    table = doc.tables[0]
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    for name, desc, hours, sig in effort_rows:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        row.cells[2].text = hours
        row.cells[3].text = sig

    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[2].text = total_hours

    out_path = OUT_DIR / out_name
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def build_week7():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 7 (May 11\u201316, 2026)"

    accomplishments = [
        "Completed the research poster (48\u00d736 inches, dark theme, programmatic layout). "
        "Two-row design presents all seven layers: Row 1 covers Layer 1 (network topology \u2014 "
        "19:1 robustness advantage), Layer 1b (single-cell economy \u2014 PBMC division of labor), "
        "and Layer 2 (metabolic simulation \u2014 71% vs 53% GDP retention). Row 2 covers "
        "Layer 2b (FBA shadow prices as subjective value), Layer 3 (cross-species trade costs), "
        "and Layers 4\u20137 (immune repertoire, genome architecture, convergent evolution, viral "
        "communication). Built via build_poster.py with embedded figures. Exported as "
        "capstone_poster.pptx.",

        "Built the Introduction section DOCX per BME 129C formatting guidelines (Arial 11, "
        "single-spaced, figure in text box with legend at font 10). Four sections: background "
        "and motivation, knowledge gap (the machine metaphor fails), specific aims (test whether "
        "living systems are literal decentralized economies), and significance for synthetic "
        "biology. Embeds the Price System of the Cell as Figure 1 with full legend.",

        "Built the Results section DOCX per BME 129C formatting guidelines. Structured as "
        "subsections per layer, each following the rationale \u2192 brief methods \u2192 result \u2192 "
        "conclusion format. Embedded figures with legends at font 10: robustness curves, "
        "GDP retention time series, FBA shadow price heatmaps, trade cost matrices, cancer "
        "mutation frequency plots. Covers all seven layers with key statistics.",

        "Built the Methods section DOCX documenting the full computational pipeline: RegulonDB "
        "network fetch and topology analysis, CellxGene single-cell PBMC processing, agent-based "
        "metabolic simulation (distributed vs centralized), flux balance analysis with iML1515 "
        "(2,712 reactions, 1,877 metabolites), codon-based trade cost calculation via Kazusa, "
        "immune repertoire analysis (IMGT, SHM hotspots, V(D)J bias), and genome-wide "
        "mutation/tissue specificity (GTEx, ClinVar).",

        "Generated seven viral network figures (generate_viral_network_figures.py) backing "
        "the viruses-as-communication essay: ERV genome composition (8% vs 1.5% protein-coding), "
        "syncytin conservation under purifying selection, global phage network (10^31 particles), "
        "gut virome composition (>90% phages), ERV regulatory elements as gene switches, "
        "autoimmune/allergic rise inverse-correlated with infection exposure, and viral "
        "communication summary across scales.",

        "Wrote \u201cThe Disease of Centralization\u201d chapter \u2014 a standalone essay arguing that "
        "disease is not a property of the variant but what happens when the environment refuses "
        "to meet the variant where it lives. Frames cancer, autoimmune disease, and metabolic "
        "syndrome as consequences of centralizing systems that evolved to be distributed.",
    ]

    challenges = [
        "Paper scope management. With all seven layers, the poster, three formatted sections, "
        "and multiple standalone essays, the project has grown substantially beyond the original "
        "3-layer proposal. Need to consolidate the written deliverables into a single coherent "
        "submission document for the final paper deadline.",

        "Sensitivity analysis on simulation parameters deferred again. The poster and formatted "
        "sections consumed the full week. The 71% vs 53% GDP result is consistent across runs "
        "but a systematic parameter sweep (agent count, perturbation magnitude, feedback gain) "
        "would strengthen the claim. Targeting Week 8.",

        "Final presentation structure still open. Seven layers is too many for a short talk. "
        "Need to select 3\u20134 layers that tell the tightest story and use the Price System "
        "figure as the anchor. Candidates: Layer 1 (topology), Layer 2 (metabolic GDP), "
        "Layer 2b (FBA shadow prices), and Panel D (cancer breaks the price system).",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Built research poster (48\u00d736, dark theme, all 7 layers). Programmatic layout "
         "via build_poster.py with embedded figures and key statistics.",
         "5 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built Introduction, Results, and Methods section DOCXs per BME 129C formatting "
         "guidelines (Arial 11, single-spaced, figures in text boxes with legends).",
         "5 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Generated 7 viral network figures. Wrote Disease of Centralization chapter. "
         "Updated run_all.py pipeline.",
         "3 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Week 7 progress report. Reviewed poster layout and formatted sections "
         "for consistency.",
         "1 hr", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "14 hrs",
                 "WeeklyUpdateReport_Week7.docx")


def build_week8():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 8 (May 18\u201323, 2026)"

    accomplishments = [
        "Expanded \u201cThe Living Architecture: Distributed Knowledge as the Design Language "
        "of Life\u201d into the comprehensive capstone synthesis paper (~98,000 characters, "
        "16 annotated figures, 45+ citations). Integrates all seven layers of evidence, "
        "six design principles for synthetic biology, the vaccines-as-central-planners "
        "application, and theological framework. Exported as a 6.8 MB DOCX with all "
        "figures embedded (the_living_architecture.docx).",

        "Built the Distributed Knowledge paper as a parallel academic version "
        "(distributed_knowledge_paper.md \u2192 distributed_knowledge_paper.docx). Same "
        "seven-layer evidence structure but written in standard academic voice with "
        "formal abstract, structured sections, and quantitative claims front-loaded. "
        "This serves as the submittable version alongside the full Living Architecture.",

        "Created two presentation decks from the Distributed Knowledge paper: a full "
        "22-slide deck (distributed_knowledge_slides.pptx) and a condensed 3-minute "
        "version (distributed_knowledge_3min.pptx). Both use the dark theme with green/"
        "gold accent palette. The 3-minute version focuses on three key results: 19:1 "
        "robustness ratio, 71% vs 53% GDP retention, and the Price System synthesis.",

        "Built the full figure presentation deck (capstone_figures_v2.pptx) \u2014 every "
        "figure from the capstone on its own slide with annotations, key statistics, "
        "and speaker notes. Designed for the defense Q&A: any question about a specific "
        "result can be answered by navigating to that figure\u2019s slide.",

        "Created title/abstract page builder (build_title_abstract.py) and intuition "
        "abstract builder (build_intuition_abstract.py) for formal submission formatting. "
        "Built capstone_paper.pdf as the final compiled paper (2.4 MB). Generated the "
        "capstone summary visual (capstone_summary_visual.pdf, 4.6 MB) as a single-page "
        "overview of all results.",

        "Built the full PDF generation pipeline (build_pdf.py) and QA document exports "
        "(build_qa_pdfs.py \u2192 capstone_QA_full.pdf, capstone_QA_bullets.pdf). The QA "
        "documents compile all six oral exam questions with detailed answers and bullet-"
        "point summaries for quick reference during the defense.",
    ]

    challenges = [
        "Final presentation timing. Seven layers plus design principles plus application "
        "is too much for a short defense talk. Decision made: anchor on the Price System "
        "figure (4 panels), support with Layer 1 (topology) and Layer 2 (metabolic GDP), "
        "reference remaining layers briefly. The figure deck is the backup for deep dives "
        "during Q&A.",

        "Paper length management. The Living Architecture is ~98K characters \u2014 comprehensive "
        "but far beyond a standard capstone paper length. The Distributed Knowledge version "
        "serves as the formal submission (~62K characters). Both exist as deliverables but "
        "the DK paper is the one formatted for academic review.",

        "Sensitivity analysis on agent-based simulation parameters remains a known limitation. "
        "The 71% vs 53% result is consistent but not formally swept. Acknowledged in the "
        "paper\u2019s limitations section rather than deferred further \u2014 the qualitative result "
        "(distributed outperforms centralized under perturbation) is robust; the exact "
        "magnitude depends on parameters.",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Expanded The Living Architecture to 98K chars. Built Distributed Knowledge "
         "paper DOCX. Two versions: full synthesis and academic submission.",
         "5 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built DK presentation decks (22-slide full + 3-minute condensed). "
         "Built figure presentation deck (capstone_figures_v2.pptx).",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built title/abstract, intuition abstract, PDF pipeline, QA document "
         "exports, and capstone summary visual.",
         "3 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Week 8 progress report. Final presentation preparation and "
         "slide ordering for capstone defense.",
         "2 hrs", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "14 hrs",
                 "WeeklyUpdateReport_Week8.docx")


if __name__ == "__main__":
    build_week7()
    build_week8()
