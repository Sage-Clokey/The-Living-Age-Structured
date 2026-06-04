"""
Build WeeklyUpdateReport DOCX for Week 8.
Follows the template format from WeeklyUpdateReportTemplate.docx.

Usage: python paper/builders/build_weekly_report_week8.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

TEMPLATE = Path(__file__).resolve().parent.parent / "WeeklyUpdateReportTemplate.docx"
OUT_DIR = Path(__file__).resolve().parent.parent / "weekly_reports"


def build_report(header, accomplishments, challenges, effort_rows, total_hours, out_name):
    doc = Document(str(TEMPLATE))

    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(header)
    run.font.size = Pt(11)

    while len(doc.paragraphs) > 4:
        p = doc.paragraphs[4]
        p._element.getparent().remove(p._element)

    for i, acc in enumerate(accomplishments, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Accomplishment {i}: {acc}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Challenges:")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()

    for i, ch in enumerate(challenges, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Challenge {i}: {ch}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    table = doc.tables[0]
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    for name, desc, hours, sig in effort_rows:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        row.cells[2].text = hours
        row.cells[3].text = sig

    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[2].text = total_hours

    out_path = OUT_DIR / out_name
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def build_week8():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 8 (May 18\u201323, 2026)"

    accomplishments = [
        "Completed the research poster (48\u00d736 inches, dark theme). The poster presents all "
        "seven layers of evidence in a two-row layout: Layer 1 (network topology), Layer 1b "
        "(single-cell economy), Layer 2 (metabolic simulation), Layer 2b (FBA/price system), "
        "Layer 3 (cross-species trade), and Layers 4\u20137 (immune, genome, convergent evolution, "
        "viral communication). Built programmatically via build_poster.py with embedded figures "
        "and key statistics. Exported as capstone_poster.pptx.",

        "Built the Introduction section DOCX per BME 129C formatting guidelines (Arial 11, "
        "single-spaced, figure in text box with legend at font 10). Four sections per course "
        "requirements: background, knowledge gap, specific aims, and significance. Embeds the "
        "Price System of the Cell figure as Figure 1 with full legend.",

        "Built the Results section DOCX per BME 129C formatting guidelines. Covers all major "
        "results across the seven layers with embedded figures and legends: network robustness "
        "(19:1 advantage), metabolic GDP retention (71% vs 53%), FBA shadow prices as subjective "
        "value, cross-species trade costs scaling with evolutionary distance, immune repertoire "
        "convergence, and cancer mutation mapping to price system components.",

        "Built the Methods section DOCX documenting the full computational pipeline: RegulonDB "
        "network analysis, single-cell PBMC transcriptomics (CellxGene), agent-based metabolic "
        "simulation, flux balance analysis (iML1515 genome-scale model), codon-based trade cost "
        "calculation (Kazusa), immune repertoire analysis (IMGT/SHM hotspots/V(D)J bias), and "
        "genome-wide mutation/tissue specificity analysis (GTEx, ClinVar).",

        "Expanded \u201cThe Living Architecture\u201d into a comprehensive synthesis document "
        "(~98,000 characters, 16 annotated figures, 45+ citations). This is the full capstone "
        "paper integrating all seven layers, design principles for synthetic biology, the "
        "vaccines-as-central-planners application, and theological framework. Exported as a "
        "6.8 MB DOCX with all figures embedded.",

        "Built the full figure presentation deck (capstone_figures_v2.pptx) \u2014 every figure "
        "from the capstone on its own slide with annotations, key statistics, and speaker notes. "
        "Also created the title/abstract page builder and intuition abstract for submission.",
    ]

    challenges = [
        "Final presentation timing. The seven layers of evidence plus design principles plus "
        "application (vaccines/cancer) is a lot of ground to cover in a short talk. Need to "
        "decide which 3\u20134 layers to present in depth and which to reference briefly. The "
        "Price System figure is the natural anchor \u2014 everything else supports it.",

        "Sensitivity analysis on agent-based simulation parameters still needs to be tightened. "
        "The 71% vs 53% GDP retention result is robust across runs but the specific parameters "
        "(number of agents, perturbation magnitude, feedback gain) should be swept systematically "
        "to show the result holds across a range of conditions.",

        "Paper length management. The Living Architecture is now ~98K characters \u2014 comprehensive "
        "but too long for a single capstone submission. Need to extract the core argument "
        "(Layers 1\u20133 + Price System + one application) into a concise 15\u201320 page version "
        "while keeping the full document as a reference.",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Built research poster (48x36, dark theme, all 7 layers). "
         "Programmatic layout via build_poster.py with embedded figures.",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built Introduction, Results, and Methods section DOCXs per BME 129C "
         "formatting guidelines (Arial 11, single-spaced, figures in text boxes).",
         "5 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Expanded The Living Architecture synthesis paper to 98K characters. "
         "Built figure presentation deck and title/abstract page.",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Week 8 progress report. Final presentation preparation and "
         "slide ordering for capstone defense.",
         "2 hrs", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "15 hrs",
                 "WeeklyUpdateReport_Week8.docx")


if __name__ == "__main__":
    build_week8()
