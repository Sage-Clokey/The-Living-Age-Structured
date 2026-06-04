"""
Build Introduction section DOCX for BME 129C Capstone.
Arial 11, single-spaced. Figure 1 in a text box with legend (font 10).
Four sections per BME129C_slides_050426.pdf guidelines.

Usage: python paper/build_introduction.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as PP_ALIGN
from docx.oxml.ns import qn
import lxml.etree as etree

OUT = Path(__file__).resolve().parent.parent / "introduction_section.docx"
FIGURE = Path(__file__).resolve().parent.parent / "figures" / "price_system_of_the_cell.png"

FONT_NAME = "Arial"
FONT_SIZE = Pt(11)
FIG_FONT_SIZE = Pt(10)


def set_single_spacing(paragraph):
    """Set paragraph to single spacing with no extra space before/after."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0


def add_paragraph(doc, text, bold=False, font_size=FONT_SIZE, alignment=PP_ALIGN.LEFT,
                  space_after=Pt(6)):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = 1.0
    pf.alignment = alignment

    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold
    return p


def add_rich_paragraph(doc, segments, space_after=Pt(6)):
    """Add paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = space_after
    pf.line_spacing = 1.0

    for seg in segments:
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.bold = bold
        run.font.italic = italic
    return p


def build():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    add_paragraph(doc, "Introduction", bold=True, font_size=Pt(12),
                  alignment=PP_ALIGN.CENTER, space_after=Pt(12))

    # =========================================================================
    # SECTION 1: Context / Importance
    # =========================================================================
    add_rich_paragraph(doc, [
        ("There are two types of order in nature: centralized order, which requires uniformity, and decentralized order, which requires diversity. Living systems are decentralized. Every cell acts on local knowledge of its own conditions and signals its neighbors voluntarily, yet the result is coordinated organization across trillions of cells without any master controller [1,2]. Synthetic biology, however, overwhelmingly treats living systems as machines\u2014specifying every promoter, every codon, every ribosome binding site from above, as though the cell were a factory floor awaiting instructions [3,4]. This is the machine metaphor applied to molecules: it reduces the system to parts, and in doing so makes what was once robust brittle. Engineered genetic circuits fail at high rates precisely because this reductionist approach destroys the distributed knowledge embedded in the system it attempts to control [5,6]. ", False, False),
        ("The persistent failure of centrally designed biological constructs suggests that the problem is not in the details of circuit design but in the paradigm itself\u2014that living things are not machines, and engineering them as if they were misrepresents how biological order actually works.", True, False),
    ])

    # =========================================================================
    # SECTION 2: Background / Overview of field
    # =========================================================================
    add_rich_paragraph(doc, [
        ("In biology there are no central banks and no legal tender laws\u2014there is only voluntary trade. Cells exchange molecular signals, but they react to those signals by choice: a cytokine does not command a T cell to activate any more than a price tag forces a customer to buy [7]. The diversity of life is not random difference\u2014it is ordered diversity based on local knowledge of time and place that no central planner can possess [8]. Network analyses have confirmed this architecture at the molecular level. Gene regulatory networks in ", False, False),
        ("Escherichia coli", False, True),
        (" exhibit scale-free degree distributions with no single master regulator, and protein-protein interaction networks in both ", False, False),
        ("E. coli", False, True),
        (" and ", False, False),
        ("Saccharomyces cerevisiae", False, True),
        (" show distributed hub structures that tolerate removal of substantial fractions of nodes before losing connectivity [1,9]. Feed-forward loops\u2014the most over-represented network motif in transcriptional regulation\u2014function as local signal-processing modules, enabling rapid coordination without routing information through a central node [10,11]. At the cellular level, single-cell RNA sequencing has shown that immune cell types specialize into complementary roles through local cytokine signaling, without any master cell assigning fates [12,13].", False, False),
    ])

    add_rich_paragraph(doc, [
        ("The distributed knowledge that coordinates living systems is dictated by prices\u2014the ratios between voluntary exchanges at every node. Austrian economics provides the formal theory for this: Hayek demonstrated that the knowledge required for rational coordination is dispersed among individual agents as tacit, local knowledge that cannot be aggregated by any central authority, and that prices are the distributed signal that transmits this knowledge without concentrating it [8,14]. Mises showed that without prices emerging from voluntary exchange, a central planner has no mechanism to determine rational resource allocation [15]. In the cell, metabolite concentrations, receptor occupancies, and energy ratios function exactly as these distributed price signals\u2014each node making educated guesses and adjusting to local conditions, and in doing so the whole is ordered. If DNA is a language, the cell speaks it and can write in that language. Bioinformatics is the study of this distributed knowledge written into the sequence of DNA [16].", False, False),
    ])

    # =========================================================================
    # SECTION 3: Current understanding AND gaps
    # =========================================================================
    add_rich_paragraph(doc, [
        ("Previous work has characterized individual properties of biological networks\u2014scale-free topology [1], robustness to random failure [9], motif enrichment [10]\u2014and has modeled metabolic systems using constraint-based optimization [17]. However, these analyses treat each property in isolation rather than asking what unifying principle connects network architecture, metabolic coordination, cellular specialization, and cross-species gene compatibility. Systems biology has largely adopted the machine framework, modeling cells as circuits to be optimized rather than economies to be understood [18,19]. This is the same category error that central planning makes in economics: treating distributed knowing as if it could be centrally held, treating knowers as if they were merely data. Synthetic biology design tools\u2014codon optimization algorithms, promoter libraries, genetic circuit CAD platforms\u2014implicitly assume the engineer can specify optimal parameters from outside the system [4,20]. The concept that biological coordination operates through a genuine price system has been noted qualitatively [21] but never tested quantitatively across multiple scales. ", False, False),
        ("What is missing is a systematic, multi-scale analysis that tests whether biological systems are structured as decentralized economies\u2014and whether reducing them to centrally planned machines explains why engineered constructs are fragile while natural systems are robust.", True, False),
    ])

    # =========================================================================
    # SECTION 4: Summary of your study
    # =========================================================================
    add_rich_paragraph(doc, [
        ("Here, we report a computational analysis testing whether biological coordination operates through distributed knowledge across three scales of organization. First, we compare the network topology of five biological networks (gene regulatory, metabolic, and protein-protein interaction networks in ", False, False),
        ("E. coli", False, True),
        (" and ", False, False),
        ("S. cerevisiae", False, True),
        (") against five synthetic architectures representing centralized, random, and structured alternatives, measuring degree distribution, betweenness centrality, and robustness under node removal. Second, we model thirteen metabolic pathways as economic agents under distributed versus centralized resource allocation, testing whether local feedback alone\u2014the biological price system\u2014achieves superior robustness despite lacking global information. Third, we analyze cross-species gene transferability across eight organisms from four kingdoms as a trade network, using codon usage distance as a measure of exchange friction. We synthesize these results into a unified framework\u2014the price system of the cell\u2014showing that intracellular metabolite ratios function as costs of capital, intercellular signals function as market prices, and signaling integrators such as mTOR function as molecular entrepreneurs reading distributed price information (Fig. 1). ", False, False),
        ("This analysis provides quantitative evidence that living things are not machines\u2014that biological order emerges from distributed knowledge at every node\u2014and that synthetic biology should design with this living architecture rather than against it.", True, False),
    ])

    # =========================================================================
    # FIGURE 1 in a text box
    # =========================================================================
    # Add some space before figure
    add_paragraph(doc, "", space_after=Pt(6))

    # Add the figure as an inline image
    if FIGURE.exists():
        # Add figure
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = PP_ALIGN.CENTER
        run = p.add_run()
        run.add_picture(str(FIGURE), width=Inches(6.5))

        # Figure legend
        legend = doc.add_paragraph()
        pf = legend.paragraph_format
        pf.space_before = Pt(4)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        bold_run = legend.add_run("Figure 1. The price system of the cell: life as a decentralized economy. ")
        bold_run.font.name = FONT_NAME
        bold_run.font.size = FIG_FONT_SIZE
        bold_run.font.bold = True

        legend_text = (
            "(A) The cell reads prices at three tiers. Intracellular metabolite ratios "
            "(ATP/ADP, NAD+/NADH, AMP/ATP) function as the cost of capital\u2014signals of the "
            "cell\u2019s own energetic state. Intercellular signals (cytokines, morphogens, growth "
            "factors, oxygen tension) function as market prices\u2014tissue-level information about "
            "supply and demand. mTOR integrates all price inputs and makes a grow-or-conserve "
            "decision without central instruction. "
            "(B) Shadow prices from flux balance analysis of the E. coli genome-scale model "
            "(iML1515; 2,712 reactions) under four environmental conditions. The same metabolite "
            "has different value depending on context\u2014NADH is cheap on glucose but expensive on "
            "acetate; oxygen is near-zero aerobically but spikes under anaerobic conditions. This "
            "is Menger\u2019s subjective value theory [13] measured in a metabolic model. "
            "(C) Metabolite pool concentrations over 200 simulation time steps in the distributed "
            "allocation regime. Early oscillations represent the price discovery phase; late "
            "convergence represents equilibrium found through local feedback alone, with no "
            "central planner. "
            "(D) Mutation frequency of key signaling genes across cancer types (COSMIC database). "
            "Each dot is one cancer type. The most frequently mutated genes\u2014TP53, PIK3CA, PTEN, "
            "mTOR, EGFR\u2014all encode components of the cellular price system. The variation across "
            "cancer types (e.g., TP53: 96% in ovarian, 1% in thyroid) demonstrates that different "
            "tissues depend on different price components. Cancer is a price system disease: the "
            "cell loses its ability to read distributed signals correctly."
        )
        text_run = legend.add_run(legend_text)
        text_run.font.name = FONT_NAME
        text_run.font.size = FIG_FONT_SIZE

    # =========================================================================
    # REFERENCES
    # =========================================================================
    add_paragraph(doc, "", space_after=Pt(6))
    add_paragraph(doc, "References", bold=True, space_after=Pt(6))

    refs = [
        "[1] Barab\u00e1si, A.-L. & Oltvai, Z.N. Network biology: understanding the cell\u2019s functional organization. Nature Reviews Genetics 5, 101\u2013113 (2004).",
        "[2] Kitano, H. Biological robustness. Nature Reviews Genetics 5, 826\u2013837 (2004).",
        "[3] Kwok, R. Five hard truths for synthetic biology. Nature 463, 288\u2013290 (2010).",
        "[4] Cameron, D.E., Bashor, C.J. & Collins, J.J. A brief history of synthetic biology. Nature Reviews Microbiology 12, 381\u2013390 (2014).",
        "[5] Purnick, P.E.M. & Weiss, R. The second wave of synthetic biology: from modules to systems. Nature Reviews Molecular Cell Biology 10, 410\u2013422 (2009).",
        "[6] Cardinale, S. & Arkin, A.P. Contextualizing context for synthetic biology \u2014 identifying causes of failure of synthetic biological systems. Biotechnology Journal 7, 856\u2013866 (2012).",
        "[7] Ramilowski, J.A. et al. A draft network of ligand\u2013receptor-mediated multicellular signalling in human. Nature Communications 6, 7866 (2015).",
        "[8] Hayek, F.A. The use of knowledge in society. American Economic Review 35, 519\u2013530 (1945).",
        "[9] Albert, R., Jeong, H. & Barab\u00e1si, A.-L. Error and attack tolerance of complex networks. Nature 406, 378\u2013382 (2000).",
        "[10] Shen-Orr, S.S., Milo, R., Mangan, S. & Alon, U. Network motifs in the transcriptional regulation network of Escherichia coli. Nature Genetics 31, 64\u201368 (2002).",
        "[11] Mangan, S. & Alon, U. Structure and function of the feed-forward loop network motif. Proceedings of the National Academy of Sciences 100, 11980\u201311985 (2003).",
        "[12] Zheng, G.X.Y. et al. Massively parallel digital transcriptional profiling of single cells. Nature Communications 8, 14049 (2017).",
        "[13] Wolf, F.A., Angerer, P. & Theis, F.J. SCANPY: large-scale single-cell gene expression data analysis. Genome Biology 19, 15 (2018).",
        "[14] Hayek, F.A. The pretence of knowledge. Nobel Memorial Lecture, December 11, 1974. American Economic Review 79, 3\u20137 (1989).",
        "[15] Mises, L. von. Economic calculation in the socialist commonwealth. In Collectivist Economic Planning (ed. Hayek, F.A.) 87\u2013130 (Routledge, 1920/1935).",
        "[16] Shapiro, J.A. How life changes itself: the Read-Write (RW) genome. Physics of Life Reviews 10, 287\u2013323 (2013).",
        "[17] Orth, J.D., Thiele, I. & Palsson, B.\u00d8. What is flux balance analysis? Nature Biotechnology 28, 245\u2013248 (2010).",
        "[18] Kitano, H. Systems biology: a brief overview. Science 295, 1662\u20131664 (2002).",
        "[19] Alon, U. An Introduction to Systems Biology: Design Principles of Biological Circuits. Chapman & Hall/CRC (2007).",
        "[20] Salis, H.M., Mirsky, E.A. & Voigt, C.A. Automated design of synthetic ribosome binding sites to control protein expression. Nature Biotechnology 27, 946\u2013950 (2009).",
        "[21] Krakauer, D.C. & Flack, J.C. Better living through physics. Nature 467, 661 (2010).",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(1)
        pf.line_spacing = 1.0
        run = p.add_run(ref)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
