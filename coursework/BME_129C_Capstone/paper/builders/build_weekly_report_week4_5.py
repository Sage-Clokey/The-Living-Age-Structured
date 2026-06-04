"""
Build WeeklyUpdateReport DOCX files for Weeks 4 and 5.
Follows the template format from WeeklyUpdateReportTemplate.docx.

Usage: python paper/build_weekly_report_week4_5.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt
from copy import deepcopy

TEMPLATE = Path(__file__).resolve().parent.parent / "WeeklyUpdateReportTemplate.docx"
OUT_DIR = Path(__file__).resolve().parent


def build_report(header, accomplishments, challenges, effort_rows, total_hours, out_name):
    doc = Document(str(TEMPLATE))

    # Paragraph 0: header
    doc.paragraphs[0].clear()
    run = doc.paragraphs[0].add_run(header)
    run.font.size = Pt(11)

    # Paragraph 2: "Summary of key accomplishments:" — already there

    # Remove template placeholder accomplishments (paragraphs 4, 5)
    # and challenges (9, 10). We'll clear all content paragraphs and rewrite.
    # Strategy: clear paragraphs 4+ and rebuild.
    # Easier: just set text on existing paragraphs and add new ones as needed.

    # Collect all paragraphs to remove (indices 4 onward, before the table)
    # We'll rebuild from scratch after paragraph 3.

    # Delete paragraphs 4 through end
    while len(doc.paragraphs) > 4:
        p = doc.paragraphs[4]
        p._element.getparent().remove(p._element)

    # Now we have paragraphs 0-3: header, blank, "Summary of key accomplishments:", blank
    # Add accomplishments
    for i, acc in enumerate(accomplishments, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Accomplishment {i}: {acc}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        # Add blank line after
        doc.add_paragraph()

    # Challenges header
    p = doc.add_paragraph()
    run = p.add_run("Challenges:")
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    doc.add_paragraph()

    for i, ch in enumerate(challenges, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"Challenge {i}: {ch}")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        doc.add_paragraph()

    # Fill the effort table
    table = doc.tables[0]

    # Remove existing data rows (keep header row 0)
    while len(table.rows) > 1:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)

    for name, desc, hours, sig in effort_rows:
        row = table.add_row()
        row.cells[0].text = name
        row.cells[1].text = desc
        row.cells[2].text = hours
        row.cells[3].text = sig

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    total_row.cells[2].text = total_hours

    out_path = OUT_DIR / out_name
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


def build_week4():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 4 (April 14\u201318, 2026)"

    accomplishments = [
        "Integrated Hayek\u2019s The Fatal Conceit (1988) across the capstone paper and slide decks. "
        "Key passage: \u2018The curious task of economics is to demonstrate to men how little they "
        "really know about what they imagine they can design.\u2019 Applied directly to synthetic "
        "biology: redesigning evolved order destroys information the designer cannot possess.",

        "Wrote Section 4.2: Design Principles \u2014 Engineering Economies, Not Machines. Two concrete "
        "architectures for sagent engineering: (1) single organism designed FOR evolution (feedback "
        "promoters, codon harmonization, shadow-price-informed expression), and (2) microbial "
        "consortium with division of labor mirroring the PBMC data from Layer 1b.",

        "Designed Section 4.7: Experimental Spread \u2014 6 Violacein Pathway Experiments. Each "
        "experiment tests one Austrian economic prediction with wet-lab data. Violacein (vioABCDE "
        "from C. violaceum) is a purple pigment quantifiable at 575nm. Experiments: (1) Fixed vs "
        "Feedback Promoters (Hayek), (2) Star vs Distributed Control (Mises), (3) Single Strain "
        "vs Consortium (Menger), (4) Codon Optimization vs Harmonization (Layer 3), (5) Adaptive "
        "Evolution over 100 generations (Kirzner), (6) Perturbation Gauntlet (Rothbard).",

        "Overhauled all figure code with data-point-forward visualization. New principle: variation "
        "IS the data \u2014 averages hide what matters. Bar charts replaced with semi-transparent "
        "bars plus scatter overlay, robustness curves show individual trial lines, hub erosion "
        "uses per-hub jitter plots with median lines. Updated 6 analysis files, regenerated all "
        "21 figures through the full pipeline.",

        "Fixed price_signals bug in single_cell_economy.py and regenerated all Layer 1b figures. "
        "Built progress report slide deck (Week 4, 7 slides) with speaker notes.",
    ]

    challenges = [
        "Need to add formal statistical validation across all layers. Mann-Whitney U for "
        "distributed vs centralized GDP, bootstrap confidence intervals on robustness metrics, "
        "KS test for power-law degree distribution fit. Planned for Week 5.",

        "Experimental spread needs advisor feedback on feasibility. The 6 violacein experiments "
        "span promoter engineering, consortium design, directed evolution, and stress testing \u2014 "
        "need Dubois input on what\u2019s achievable with available wet-lab resources. Prioritize "
        "experiments 1\u20133 if resources are limited.",

        "Paper currently at 3 layers of evidence. The immune system and genome-wide analyses "
        "are natural extensions but would require significant new analysis pipelines. Deciding "
        "whether to deepen existing layers or broaden to new ones.",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Integrated Fatal Conceit across paper. Wrote Design Principles section (two architectures). "
         "Designed experimental spread (6 violacein experiments).",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Overhauled all figure code: data-point-forward visualization across 6 analysis files. "
         "Regenerated all 21 figures. Fixed price_signals bug.",
         "3 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built Week 4 progress report slides (7 slides with speaker notes). "
         "Paper editing and formatting.",
         "1.5 hrs", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "8.5 hrs",
                 "WeeklyUpdateReport_Week4.docx")


def build_week5():
    header = "Student: Sage Clokey (Individual Project) \u2014 Week 5 (April 19\u201325, 2026)"

    accomplishments = [
        "Extended the project from 3 layers to 7 layers of biological evidence. Added Layer 4 "
        "(Immune System): SHM hotspot analysis showing 19:1 enrichment at WRC motifs, V(D)J "
        "segment usage bias with IGHV3-23 at 10\u201320x rare segments, and public clonotypes "
        "shared across unrelated individuals at rates 10^15 above random chance.",

        "Added Layer 5 (Whole Genome): CpG C>T mutation hotspots at 15\u201340x baseline, "
        "transition/transversion ratio of 2:1 (4x random expectation), and tissue-specific "
        "gene expression showing 20% of genes with tau > 0.95 (near-exclusive expression in "
        "one tissue). Added Layer 6 (Convergent Evolution): 35 documented events across kingdoms "
        "including 14 identical Prestin substitutions between bats and dolphins.",

        "Added Layer 7 (Viral Communication): entirely new analysis framing viruses as horizontal "
        "information transfer in the cellular economy. 8% of human genome is endogenous retrovirus. "
        "Syncytin (captured viral gene) enables mammalian pregnancy. 10^31 phages conduct 10^25 "
        "gene transfers per day \u2014 the first internet. Wrote standalone essay with 7 data-backed "
        "figures.",

        "Wrote \u201cThe Living Architecture\u201d \u2014 55-page synthesis paper (715 lines, 13,625 words, "
        "16 annotated figures, 45 citations). Integrates all 7 layers from network topology through "
        "viral communication, plus design principles, vaccines as application, and theological "
        "framework. Full DOCX with embedded annotated figures (6.6 MB).",

        "Built automated figure annotation pipeline (annotate_figures.py). Overlays callout boxes "
        "with arrows and key findings on all 16 core figures. Three visual styles: teal (findings), "
        "red (warnings), gold (highlights). Total figure count went from 21 to 39.",

        "Wrote \u201cVaccines as Central Planners\u201d standalone essay mapping 5 vaccine-induced "
        "distortions to Austrian economics: germinal center monopolization, original antigenic sin, "
        "tissue misallocation, feedback suppression, and immune repertoire narrowing.",
    ]

    challenges = [
        "Poster layout deferred to Week 6. The scope expansion from 3 to 7 layers consumed all "
        "available time. Built 22-slide and 3-minute presentation decks as intermediate "
        "deliverables instead.",

        "Live API pipeline run deferred. The immune and genome analyses use published empirical "
        "data directly (IMGT, GTEx, ClinVar), so the API integration for Layers 1\u20133 was less "
        "critical than extending the evidence base. Will revisit when preparing final figures.",

        "Need advisor review of the Living Architecture paper. Key question: with 7 layers of "
        "evidence, which are strongest for the final presentation? Should the final talk focus "
        "on 3\u20134 layers deeply or touch all 7?",
    ]

    effort_rows = [
        ("Sage Clokey",
         "Built immune analysis pipeline (immune_distributed_knowledge.py): SHM hotspots, "
         "V(D)J bias, public clonotypes. Built genome analysis pipeline "
         "(genome_distributed_knowledge.py): mutation hotspots, tissue specialization, "
         "convergent evolution.",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Wrote viral communication essay with 7 figures. Created viral analysis pipeline. "
         "Wrote vaccines essay.",
         "3 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Wrote The Living Architecture synthesis paper (55 pages, 715 lines). "
         "Built annotate_figures.py pipeline. Annotated all 16 core figures.",
         "4 hrs", "Sage Clokey"),
        ("Sage Clokey",
         "Built Week 5 progress report slides (8 slides with speaker notes). "
         "Built 22-slide and 3-minute presentation decks. DOCX exports.",
         "2 hrs", "Sage Clokey"),
    ]

    build_report(header, accomplishments, challenges, effort_rows, "13 hrs",
                 "WeeklyUpdateReport_Week5.docx")


if __name__ == "__main__":
    build_week4()
    build_week5()
